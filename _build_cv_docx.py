"""CV para 30X Head of Content AI — estructura espejo del CV 2026 + tono LinkedIn v4."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r"C:\Users\ASUS\OneDrive\Escritorio\Juan Barrios\Documentos{"
OUT = FOLDER + r"\CV - Juan Camilo Barrios - Head of Content AI 30X.docx"

FONT = "Calibri"
TEXT = RGBColor(0x1a, 0x1a, 0x1a)
MUTED = RGBColor(0x55, 0x55, 0x55)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

s = doc.styles["Normal"]
s.font.name = FONT
s.font.size = Pt(10.5)
s.font.color.rgb = TEXT
s.paragraph_format.space_after = Pt(3)

def run_line(text, size=10.5, bold=False, italic=False, color=TEXT, align=None, space_after=3):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def section_title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = TEXT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def bullet(text):
    # support "**Label.** rest" pattern
    p = doc.add_paragraph(style="List Bullet")
    if text.startswith("**"):
        end = text.find("**", 2)
        if end > 0:
            label = text[2:end]
            rest = text[end+2:]
            r1 = p.add_run(label)
            r1.font.bold = True
            r1.font.size = Pt(10.5)
            r2 = p.add_run(rest)
            r2.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(2)
            return
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)

def job_header(title, meta):
    run_line(title, size=11, bold=True, space_after=0)
    run_line(meta, size=9.5, color=MUTED, italic=True, space_after=3)

# =========================================
# HEADER
# =========================================
run_line("JUAN CAMILO BARRIOS", size=22, bold=True, space_after=1)
run_line("Founder & AI Content Director, Athora AI", size=12, italic=True, color=MUTED, space_after=1)
run_line("+57 314 471 5556   |   barrioschacon.juancamilo@gmail.com   |   Bogotá, Colombia", size=10, color=MUTED, space_after=6)

# =========================================
# PROFILE
# =========================================
section_title("PROFILE")
run_line(
    "Founder of Athora AI, an AI content agency based in Colombia. We produce what a normal "
    "agency takes three months to produce, in twenty four hours. No cameras, no actors, no "
    "studio. Just a stack that works: Seedance 2.0, Higgsfield, Kling 3.0, Nano Banana, Seed "
    "Dream 2.0, Claude, and n8n, running behind a pipeline I built and refined over the last year."
)
run_line(
    "I lead the full ecosystem: brand strategy, content direction, production, automation, and "
    "client work. I write the strategy, direct the shots, wire the workflows, and close the "
    "deals. I also build internal tools and specialist agents when the workflow asks for them, "
    "which is often."
)
run_line(
    "Before Athora I spent close to a decade inside customer operations and high-end kitchens: "
    "Booking.com, Capital One, Bet365, El Chato, Arzak. That background shaped how I build. "
    "Fast, systems-first, obsessive about the detail, and aware that every deliverable lands "
    "on a real person's screen.",
    space_after=6,
)

# =========================================
# CORE EXPERTISE
# =========================================
section_title("CORE EXPERTISE")
bullet("**AI Content Direction.** Cinematic video, UGC, digital avatars with character lock and voice clone. Seedance 2.0, Higgsfield, Kling 3.0, Seed Dream 2.0, Nano Banana.")
bullet("**Editorial criterion for C-suite.** I know the difference between executive content and generic feed filler. Brand pillars, positioning, voice guidelines, content calendars built for audiences that don't forgive a misstep.")
bullet("**AI-first systems, not AI-as-a-topic.** I build workflows and agents, not just prompts. 15 Seedance specialist skills, a Billion-Dollar Board (eight simulated experts), a sales-executor skill, Notion-via-API as standard.")
bullet("**Workflow Automation.** n8n, Claude API, Python, Apify, Scrapling, ElevenLabs, Google MCPs. If it can become an internal app, I build it.")
bullet("**Rapid Web Development.** Next.js, React, HTML and CSS vanilla, AI-assisted code. Brief to live site in under 48 hours (Paseo, Open Carrusel).")
bullet("**Lead Generation & B2B Outreach.** Scrapling, Apify, custom scrapers, shared CRMs. Frameworks: Kennedy cold email 11-step, Cardone 5-touch, Hormozi Grand Slam Offer.")
bullet("**Creative team management.** Athora AI: partner, freelancers, pipeline ownership end to end.")
bullet("**Bilingual client-facing.** English and Spanish, C-level calls, strategy sessions, proposal writing.")

# =========================================
# WORK EXPERIENCE
# =========================================
section_title("WORK EXPERIENCE")

job_header("Founder & AI Content Director", "Athora AI   |   Bogotá, Colombia   |   September 2024 – Present")
bullet("Founded and lead an AI content agency serving SMBs and personal brands across real estate, aesthetic and dental clinics, hospitality, luxury chocolate, corporate events, and creator brands. No client names in writing; case studies shared under NDA.")
bullet("Built a proprietary production stack combining Seedance 2.0, Higgsfield, Kling 3.0, Nano Banana, and Seed Dream 2.0. The pipeline ships cinematic commercial video, UGC, and digital avatar content at a fraction of the cost and timeline of traditional post-production.")
bullet("Run my own ultra-realistic digital twin for Athora's marketing, posting two to three organic Instagram pieces per day while I build the agency in public. The twin is the proof of concept I sell.")
bullet("Designed a content operating system used across clients: weekly content calendar of 12 pieces (40% viral / 35% educational / 25% service), one-day-shoot-for-seven-days-of-content model, brand-pillar knowledge base in Obsidian.")
bullet("Led editorial direction and full site build for Paseo (AI-powered virtual tours for realtors). Brand identity (tungsten and projector-booth palette, editorial typography), HTML and CSS vanilla implementation, cinematic hero, CSS 3D POV engine. Brief to deployed in under 48 hours.")
bullet("Delivered Smile Chocolatería brand system: six-pillar identity, seven-document knowledge base, weekly content plan of 12 reels, 7 stories, 1 carousel. Zero comparisons to other brands, ultra-realism across every piece, logo and slogan on every asset.")
bullet("Built a B2B lead generation system using Apify, Scrapling and Python. Captured and qualified 100+ Colombian fintechs and startups in an afternoon for a shared CRM (Elevate Seven).")
bullet("Ran a Madrid real estate outreach campaign in April 2026: 15 fully personalized emails with a tailored marketing video attached, CRM updated contact by contact, Kennedy 11-step follow-up and Cardone 5-touch sequence.")
bullet("Built a Grand Slam Offer for Spain real estate (\"Propiedad Cinemática™\", $3.5K/month, $27.5K perceived stack, 15 visits in 90 days guarantee) and a full Zoho SOA proposal for corporate events (Carlos Moya), including a Notion Implementation Guide built via API in a single session (266 blocks, 11 skills, runbook).")
bullet("Designed and built internal tooling: 15 Seedance specialist skills in Claude Code (cinematic, real estate, food and beverage, fashion, social hook, brand story, product 360, ecommerce ad, and more), a Billion-Dollar Board agent with eight simulated experts, a billion-deal sales executor, Open Carrusel (Next.js 16 + React 19) for Instagram carousels, Google MCPs integrated end to end.")
bullet("Close clients end to end: first outreach, discovery, proposal, strategy session, delivery, post-delivery review.")

job_header("Verification Advisor", "Bet365   |   Bogotá, Colombia   |   April 2024 – December 2025")
bullet("Validated client identity documents at volume, with strict adherence to KYC protocols and internal review standards across multiple jurisdictions.")
bullet("Flagged identity theft attempts, underage registrations and gambling-risk indicators during document review, escalating cases through the right compliance channels.")
bullet("Worked across multiple country formats and prioritized queues to keep pending verifications within SLA.")

job_header("Customer Services Advisor", "Booking.com (via 24-7 Intouch)   |   Bogotá, Colombia   |   June 2022 – July 2023")
bullet("Handled ~1,400 reservations per month with 92%+ first-contact resolution. Top 10% of the team in NPS.")
bullet("Resolved escalations in English and Spanish, negotiating directly with hotels when the standard playbook did not fit.")
bullet("Built repeat relationships by matching recommendations to real client context (budget, trip purpose, preferences) instead of running a script.")

job_header("Customer Experience Associate", "Capital One (via Sutherland)   |   Bogotá, Colombia   |   June – December 2021")
bullet("Multichannel customer service (phone, email, live chat) for a US financial services client. Resolved account and product inquiries end to end.")
bullet("Maintained clean interaction records to support reliable follow-up and audit trail across the customer lifecycle.")
bullet("Surfaced recurring friction points to internal teams and contributed to iterative updates of support scripts and workflow.")

job_header("Sous Chef", "El Chato + Arzak (San Sebastián, Spain)   |   Bogotá and Spain   |   January 2020 – December 2020")
bullet("High-end kitchens: El Chato (Latin America's 50 Best) and Arzak (3 Michelin stars). Mise en place, cold and hot stations, high-volume service.")
bullet("Supervised kitchen assistants, ran inventory, ordering and receiving.")
bullet("That year gave me a visual and operational standard I still apply when Athora produces food and hospitality content. It reads in the output.")

# =========================================
# EDUCATION
# =========================================
section_title("EDUCATION")
run_line("LaSalle College — Technological Degree in Gastronomy and Restaurant Management", bold=True, space_after=0)
run_line("September 2017 – September 2020", size=9.5, color=MUTED, italic=True, space_after=3)
run_line("RETOS School — Academic Bachelor", bold=True, space_after=0)
run_line("December 2014 – December 2015", size=9.5, color=MUTED, italic=True, space_after=3)
run_line("Self-directed (2023 – present): Claude Code (Anthropic), advanced n8n, ElevenLabs voice AI, agent architecture, LLM prompt engineering, Ramiro Cubría 8-module program on hooks, outbound and sales.", size=10.5, space_after=4)

# =========================================
# LANGUAGES
# =========================================
section_title("LANGUAGES")
run_line("English (Professional)   |   Spanish (Native)   |   Italian (Basic)")

# =========================================
# REFERENCES
# =========================================
section_title("REFERENCES")
run_line("Paola Chacón — +57 320 490 7824", space_after=1)
run_line("Arturo Barrios — +57 321 204 6244", space_after=1)
run_line("Frank Hodson — +57 318 788 3945")

doc.save(OUT)
print(f"OK -> {OUT}")
